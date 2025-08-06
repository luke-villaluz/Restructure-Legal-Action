import os
import io  # Add this import for BytesIO
from typing import List, Dict, Optional, Any
from utils.logger import logger
from config.settings import TESSERACT_PATH

class DocumentProcessor:
    """Process PDF and Word documents for text extraction with OCR fallback"""
    
    def __init__(self):
        self.logger = logger
        self.supported_extensions = ['.pdf', '.docx', '.doc']
        
    def is_document_file(self, filename: str) -> bool:
        """Check if file is a supported document type"""
        return any(filename.lower().endswith(ext) for ext in self.supported_extensions)
    
    def get_all_documents_in_folder(self, folder_path: str) -> List[str]:
        """
        Get all document files in a folder (including nested subfolders)
        
        Args:
            folder_path: Path to the folder to scan
            
        Returns:
            List of full paths to all document files
        """
        document_files = []
        
        try:
            for root, dirs, files in os.walk(folder_path):
                for file in files:
                    if self.is_document_file(file):
                        full_path = os.path.join(root, file)
                        document_files.append(full_path)
                        self.logger.info(f"Found document: {full_path}")
            
            self.logger.info(f"Found {len(document_files)} documents in {folder_path}")
            return document_files
            
        except Exception as e:
            self.logger.error(f"Error scanning folder {folder_path}: {e}")
            return []
    
    def extract_text_from_pdf(self, pdf_path: str) -> Optional[str]:
        """Extract text from a PDF file using Tesseract OCR only"""
        try:
            self.logger.info(f"Extracting text from: {pdf_path}")
            
            # Use Tesseract OCR directly
            ocr_text = self._extract_text_with_tesseract(pdf_path)
            if ocr_text:
                self.logger.info(f"Successfully extracted {len(ocr_text)} characters from PDF using Tesseract OCR")
                return ocr_text
            else:
                self.logger.error(f"Failed to extract text from PDF: {pdf_path}")
                return None
            
        except Exception as e:
            self.logger.error(f"Error extracting text from PDF {pdf_path}: {e}")
            return None
    
    def _extract_text_with_tesseract(self, pdf_path: str) -> Optional[str]:
        """
        Extract text from PDF using Tesseract OCR (much better for legal documents)
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text or None if failed
        """
        try:
            # Import OCR dependencies
            try:
                import pytesseract
                from PIL import Image, ImageEnhance, ImageFilter
                import fitz  # PyMuPDF for PDF to image conversion
            except ImportError as e:
                self.logger.error(f"OCR dependencies not installed: {e}")
                self.logger.error("Install with: pip install pytesseract Pillow PyMuPDF")
                return None
            
            # Use configurable Tesseract path
            pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
            
            self.logger.info(f"🖼️ Starting Tesseract OCR processing for: {pdf_path}")
            
            # Open PDF with PyMuPDF
            pdf_document = fitz.open(pdf_path)
            extracted_text = ""
            
            for page_num in range(len(pdf_document)):
                try:
                    # Get page
                    page = pdf_document.load_page(page_num)
                    
                    # Convert page to image with high resolution
                    mat = fitz.Matrix(3.0, 3.0)  # 3x zoom for better OCR
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Convert to PIL Image
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Enhance image for better OCR using PIL only
                    # Convert to grayscale if needed
                    if img.mode != 'L':
                        img = img.convert('L')
                    
                    # Enhance contrast
                    enhancer = ImageEnhance.Contrast(img)
                    img = enhancer.enhance(2.0)  # Increase contrast
                    
                    # Enhance sharpness
                    enhancer = ImageEnhance.Sharpness(img)
                    img = enhancer.enhance(1.5)  # Increase sharpness
                    
                    # Apply slight blur to reduce noise
                    img = img.filter(ImageFilter.GaussianBlur(radius=0.5))
                    
                    # Perform OCR with Tesseract
                    # Use PSM 6 for uniform block of text (good for legal documents)
                    custom_config = r'--oem 3 --psm 6'
                    
                    page_text = pytesseract.image_to_string(
                        img, 
                        config=custom_config,
                        lang='eng'
                    )
                    
                    if page_text.strip():
                        extracted_text += page_text.strip() + "\n"
                        self.logger.info(f"Tesseract extracted {len(page_text)} characters from page {page_num + 1}")
                    else:
                        self.logger.warning(f"Tesseract found no text on page {page_num + 1}")
                        
                except Exception as e:
                    self.logger.warning(f"Tesseract failed on page {page_num + 1}: {e}")
                    continue
            
            pdf_document.close()
            
            if extracted_text.strip():
                return extracted_text.strip()
            else:
                self.logger.warning("Tesseract extracted no text from any page")
                return None
                
        except Exception as e:
            self.logger.error(f"Tesseract processing failed: {e}")
            return None
    
    def extract_text_from_word(self, word_path: str) -> Optional[str]:
        """Extract text from a Word document (.docx and .doc)"""
        try:
            self.logger.info(f"Extracting text from Word document: {word_path}")
            
            # Handle .docx files
            if word_path.lower().endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document(word_path)
                    
                    # Extract text from paragraphs
                    text = ""
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text.strip() + "\n"
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text += cell.text.strip() + "\n"
                    
                    if text.strip():
                        self.logger.info(f"Successfully extracted {len(text)} characters from Word document (.docx)")
                        return text.strip()
                    else:
                        self.logger.warning(f"No text extracted from Word document: {word_path}")
                        return None
                        
                except ImportError:
                    self.logger.error("python-docx not installed. Install with: pip install python-docx")
                    return None
                except Exception as e:
                    self.logger.error(f"Error extracting from .docx: {e}")
                    return None
            
            # Handle .doc files
            elif word_path.lower().endswith('.doc'):
                try:
                    import win32com.client
                    import pythoncom
                    
                    # Initialize COM
                    pythoncom.CoInitialize()
                    
                    # Create Word application
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    
                    try:
                        # Open the document
                        doc = word.Documents.Open(os.path.abspath(word_path))
                        
                        # Extract text
                        text = doc.Content.Text
                        
                        # Close document
                        doc.Close()
                        
                        if text.strip():
                            self.logger.info(f"Successfully extracted {len(text)} characters from Word document (.doc)")
                            return text.strip()
                        else:
                            self.logger.warning(f"No text extracted from Word document: {word_path}")
                            return None
                            
                    finally:
                        # Quit Word application
                        word.Quit()
                        pythoncom.CoUninitialize()
                        
                except ImportError:
                    self.logger.error("pywin32 not installed. Install with: pip install pywin32")
                    return None
                except Exception as e:
                    self.logger.error(f"Error extracting from .doc: {e}")
                    return None
            
            else:
                self.logger.error(f"Unsupported Word format: {word_path}")
                return None
            
        except Exception as e:
            self.logger.error(f"Error extracting text from Word document {word_path}: {e}")
            return None
    
    def extract_text_from_document(self, document_path: str) -> Optional[str]:
        """Extract text from any supported document type"""
        try:
            if document_path.lower().endswith('.pdf'):
                return self.extract_text_from_pdf(document_path)
            elif document_path.lower().endswith(('.docx', '.doc')):
                return self.extract_text_from_word(document_path)
            else:
                self.logger.error(f"Unsupported document type: {document_path}")
                return None
                
        except Exception as e:
            self.logger.error(f"Error extracting text from document {document_path}: {e}")
            return None
    
    def extract_all_text_from_folder(self, folder_path: str) -> Dict[str, Any]:
        """
        Extract text from all documents in a folder
        
        Args:
            folder_path: Path to the folder containing documents
            
        Returns:
            Dictionary containing:
            {
                'successful_texts': Dict[str, str],  # doc_path -> text
                'failed_documents': List[str],       # list of failed doc names
                'document_stats': Dict[str, int]     # total, successful, failed counts
            }
        """
        try:
            document_files = self.get_all_documents_in_folder(folder_path)
            
            if not document_files:
                self.logger.warning(f"No documents found in {folder_path}")
                return {
                    'successful_texts': {},
                    'failed_documents': [],
                    'document_stats': {'total': 0, 'successful': 0, 'failed': 0}
                }
            
            successful_texts = {}
            failed_documents = []
            
            for doc_path in document_files:
                text = self.extract_text_from_document(doc_path)
                if text:
                    successful_texts[doc_path] = text
                else:
                    failed_documents.append(os.path.basename(doc_path))
                    self.logger.warning(f"Failed to extract text from {doc_path}")
            
            document_stats = {
                'total': len(document_files),
                'successful': len(successful_texts),
                'failed': len(failed_documents)
            }
            
            self.logger.info(f"Successfully extracted text from {len(successful_texts)} documents in {folder_path}")
            self.logger.info(f"Document stats: {document_stats}")
            
            return {
                'successful_texts': successful_texts,
                'failed_documents': failed_documents,
                'document_stats': document_stats
            }
            
        except Exception as e:
            self.logger.error(f"Error extracting text from folder {folder_path}: {e}")
            return {
                'successful_texts': {},
                'failed_documents': [],
                'document_stats': {'total': 0, 'successful': 0, 'failed': 0}
            }
    
    def combine_document_texts(self, document_texts: Dict[str, str]) -> str:
        """
        Combine text from multiple documents into a single analysis-ready text
        
        Args:
            document_texts: Dictionary mapping document paths to text
            
        Returns:
            Combined text with document separators
        """
        try:
            if not document_texts:
                return ""
            
            combined_text = ""
            
            for doc_path, text in document_texts.items():
                doc_name = os.path.basename(doc_path)
                combined_text += f"\n\n=== DOCUMENT: {doc_name} ===\n\n"
                combined_text += text
                combined_text += "\n\n"
            
            self.logger.info(f"Combined text from {len(document_texts)} documents ({len(combined_text)} total characters)")
            return combined_text.strip()
            
        except Exception as e:
            self.logger.error(f"Error combining document texts: {e}")
            return ""